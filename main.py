from fastapi import FastAPI
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import uuid
import os

app = FastAPI()

# Ensure output directory exists
OUTPUT_DIR = "generated_docs"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# -----------------------------
# Health Check (for Render)
# -----------------------------
@app.get("/health")
def health_check():
    return {"status": "ok"}


# -----------------------------
# Request Schema
# -----------------------------
class AARRequest(BaseModel):
    certification_level: str
    experience: str
    scenario_id: str
    scenario_title: str
    aar_text: str


# -----------------------------
# AAR Generator Endpoint
# -----------------------------
@app.post("/generate-aar")
def generate_aar(data: AARRequest):

    # Generate random learner ID (4-char alphanumeric)
    learner_id = uuid.uuid4().hex[:4].upper()

      # Create document
    doc = Document()

    # Header with logo
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = header_paragraph.add_run()
    run.add_picture("CCT-EDGE_TXPARENT.png", width=Inches(2.0))

    # Spacer after header
    doc.add_paragraph("")

    # Title
    doc.add_heading("After Action Report (AAR)", 0)

    # Metadata
    doc.add_paragraph(f"Learner ID: {learner_id}")
    doc.add_paragraph(f"Certification Level: {data.certification_level}")
    doc.add_paragraph(f"CCT Experience Level: {data.experience}")
    doc.add_paragraph(f"Scenario ID: {data.scenario_id}")
    doc.add_paragraph(f"Scenario Title: {data.scenario_title}")

    doc.add_paragraph("")  # spacer

    # Add AAR text
    for line in data.aar_text.split("\n"):
        doc.add_paragraph(line)

    # Save file
    filename = f"{learner_id}.docx"
    file_path = os.path.join(OUTPUT_DIR, filename)
    doc.save(file_path)
    # Return download path
    return {
        "learner_id": learner_id,
        "download_url": f"/download/{filename}"
    }


# -----------------------------
# File Download Endpoint
# -----------------------------
@app.get("/download/{filename}")
def download_file(filename: str):
    file_path = os.path.join(OUTPUT_DIR, filename)
    return FileResponse(file_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
